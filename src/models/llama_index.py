from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.core import (
    VectorStoreIndex,
    SimpleDirectoryReader,
    StorageContext,
    Settings
)
from llama_index.core.vector_stores import MetadataFilters, ExactMatchFilter
#from llama_index.core.node_parser import SentenceSplitter
#from llama_index.core.llms import ChatMessage
from s3fs import S3FileSystem
import openai
import os
import boto3
import qdrant_client
from llama_index.vector_stores.qdrant import QdrantVectorStore
from qdrant_client.http.models import Distance, VectorParams, HnswConfigDiff
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Document
#from llama_index.core.schema import TextNode
from PyPDF2 import PdfReader
from io import BytesIO,StringIO
import pandas as pd
from docx import Document as DocxDocument
from llama_index.core import PromptTemplate
from llama_index.core.prompts.prompt_type import PromptType
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core import get_response_synthesizer
from llama_index.core.retrievers import VectorIndexRetriever
from src.models.mysql_connector import execute_query,update_query,update_query_last_index, chat_bot_execute_query, chat_bot_update_query_last_index, run_query
import json

Settings.llm = OpenAI(model_name=os.environ.get('DEFAULT_MODEL'),temperature=0)
Settings.embed_model = OpenAIEmbedding(model=str(os.environ.get('EMBEDDING_MODEL')),embed_batch_size=20,max_retries=3)
Settings.node_parser = SentenceSplitter(chunk_size=int(os.environ.get('CHUNK_SIZE')), chunk_overlap=20)
Settings.num_output = 1024
Settings.context_window = 4000


class LLAMA_INDEX():
    def __init__(self) -> None: 
        self.OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
        openai.api_key = self.OPENAI_API_KEY
        self.is_storage_context = False
        self.AWS_REGION = os.environ.get('AWS_REGION')
        self.sentence_chunk_size = int(os.environ.get('CHUNK_SIZE'))      
        self.bucket_name = os.environ.get('CDN_BUCKET')
        #self.index_s3_bucket_name = os.environ.get('INDEX_S3_BUCKET_NAME') # {bucket_name}/{index_name}        
        self.S3_ACCESS_KEY = os.environ.get('S3_ACCESS_KEY')
        self.S3_SECRET_KEY = os.environ.get('S3_SECRET_KEY')
        self.QUERY_ENGINE_RESPONSE_MODE = os.environ.get('QUERY_ENGINE_RESPONSE_MODE')
        #self.s3_fs = self.initialize_s3()
        self.s3_client_obj = self.initialize_s3_client()
        self.embed_model = OpenAIEmbedding(model_name=str(os.environ.get('EMBEDDING_MODEL')),embed_batch_size=16,max_retries=3)        
        self.qdrant_url = os.environ.get('QDRANT_HOST')
        self.qdrant_client = self.initialize_qdrant()
        self.llm = OpenAI(model_name=os.environ.get('DEFAULT_MODEL'),temperature=0)        
        self.QA_PROMPT = str(os.environ.get('QA_PROMPT'))

    def initialize_s3_client(self):
        S3_ACCESS_KEY = os.environ.get('S3_ACCESS_KEY')
        S3_SECRET_KEY = os.environ.get('S3_SECRET_KEY')
        AWS_REGION = os.environ.get('AWS_REGION')
        return boto3.client('s3', aws_access_key_id=S3_ACCESS_KEY, aws_secret_access_key=S3_SECRET_KEY,region_name=AWS_REGION)

    def initialize_s3(self):
        S3_ACCESS_KEY = os.environ.get('S3_ACCESS_KEY')
        S3_SECRET_KEY = os.environ.get('S3_SECRET_KEY')
        return S3FileSystem(anon=False, key=S3_ACCESS_KEY, secret=S3_SECRET_KEY)

    def initialize_qdrant(self):
        qdrant_url = os.environ.get('QDRANT_HOST')
        return qdrant_client.QdrantClient(qdrant_url, grpc_port=6334, prefer_grpc=True)
                
    def key_exists(self, key):        
        try:
            result = self.s3_client_obj.list_objects_v2(Bucket=str(self.bucket_name), Prefix = key)
            for obj in result['Contents']:
                if obj['Key'] == key+"/":
                    return True
        except Exception as error:
            print(f"s3 key exist error : {error}")
            return False        
    
    def read_file_contents(self, file_key,file_data):
        if file_key.endswith('.csv'):
            # Read CSV data
            csv_data = file_data.decode('utf-8')
            df = pd.read_csv(StringIO(csv_data))
            text_content = df.to_string(index=False)
            return text_content

        elif file_key.endswith('.txt'):
            # Read text data
            text_data = file_data.decode('utf-8')
            return text_data

        elif file_key.endswith('.docx'):
            # Read DOCX data
            docx_data = BytesIO(file_data)
            doc = DocxDocument(docx_data)
            text_content = [para.text for para in doc.paragraphs]
            return "\n".join(text_content)

        elif file_key.endswith('.xlsx') or file_key.endswith('.xls'):
            # Read Excel data (for both .xlsx and .xls)
            xls_data = BytesIO(file_data)
            df = pd.read_excel(xls_data, engine='openpyxl' if file_key.endswith('.xlsx') else 'xlrd')
            text_content = df.to_string(index=False)
            return text_content
        
        elif file_key.endswith('.pdf'):
            pdf_reader = PdfReader(BytesIO(file_data))
            text_content = []
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())

            full_text = "\n".join(text_content)
            return full_text

        else:
            print(f"Unsupported file format {file_key}.")
            return None
    
    def createQdrantVectorIndex(self,collectionName,rag_folder_name,folder_name):
        index = None
        try:      
            #documents = self.load_documents(client_service_name)            

            documents = []
            result = self.s3_client_obj.list_objects_v2(Bucket=str(self.bucket_name), Prefix = f"{rag_folder_name}/{folder_name}")
            if 'Contents' in result:
                for obj in result['Contents']:
                    file_key = obj['Key']
                    if file_key.endswith('/'):
                        continue
                    file_response = self.s3_client_obj.get_object(Bucket=str(self.bucket_name), Key=file_key)
                    file_raw_data = file_response['Body'].read()
                    file_data = self.read_file_contents(file_key,file_raw_data)
                    #print(file_data)
                    if file_data is None:
                        continue   

                    file_name = os.path.basename(file_key)            
                    print(f"s3 docs {file_key} read successfully. file name : {file_name}")
                    
                    doc_obj = Document(text=file_data, metadata={"role": collectionName,"file_name":file_name})                 
                    documents.append(doc_obj)            

            vector_store = QdrantVectorStore(
                    client=self.qdrant_client, 
                    collection_name=str(collectionName),                    
                    embedding=self.embed_model,
                    prefer_grpc=True
                )
            print("vector_store created successfully.")
            storage_context = StorageContext.from_defaults(vector_store=vector_store)

            existing_collections = self.qdrant_client.get_collections().collections
            collection_names = [collection.name for collection in existing_collections]

            if str(collectionName) in collection_names:
                self.qdrant_client.delete_collection(collection_name=str(collectionName))

            # Create or recreate the collection in Qdrant with HNSW configuration
            self.qdrant_client.create_collection(
                collection_name=str(collectionName),
                vectors_config=VectorParams(
                    size = 1536, #openai ada embedding dimension
                    distance = Distance.COSINE,
                    hnsw_config = HnswConfigDiff(
                            ef_construct=200,  # Controls index construction speed vs. accuracy tradeoff
                            m=16,  # Number of bi-directional links created for every new element during construction
                            full_scan_threshold=10000  # Threshold for switching between HNSW and brute-force search
                        )
                    ),
            )
            print("qdrant collection created successfully.")                    
            index = VectorStoreIndex.from_documents(
                        documents,
                        embed_model=self.embed_model,
                        storage_context=storage_context
                    )
            print("vector index created successfully.")
            self.updateVectorIndexRequestStatus('C',collectionName) 
        except Exception as error:
            print("Exception in LLAMA_INDEX createQdrantVectorIndex. ",error)
            self.updateVectorIndexRequestStatus('E',collectionName)
        finally:
            return index
        
    def updateVectorIndexRequestStatus(self,status_flag,collectionName):
        try:
            #need to change
            query = "UPDATE `qdrant_collections` SET `vector_index_req`= %s WHERE collection_name = %s" 
            values = (status_flag, collectionName)
            res_obj = update_query_last_index(query, values)
            if res_obj['row_count'] > 0:
                row_count = res_obj['last_index']
                if row_count > 0:
                    print("Vector index request status flag updated.")
        except Exception as error:
            print("Exception while update VectorIndex request status. ",error)

    def insertRagResponseData(self,values):
        try:
            query = "INSERT INTO rag_transaction (sequence_id, conversation_id, rag_id, rag_score, rag_text, rag_metadata, rag_query, rag_response) VALUES (%s, %s, %s, %s, %s, %s, %s, %s);"            
            res_obj = chat_bot_update_query_last_index(query, values)
            if res_obj['row_count'] > 0:
                row_count = res_obj['last_index']
                if row_count > 0:
                    print("Rag response data stored.")
        except Exception as error:
            print("Exception while inserting rag response data. ",error)

    def getQdrantVectorIndex(self,collectionName):
        index = None
        try:            
            vector_store = QdrantVectorStore(client=self.qdrant_client, collection_name=str(collectionName))            
            index = VectorStoreIndex.from_vector_store(vector_store)
        except Exception as error:
            print("Exception in LLAMA_INDEX getQdrantVectorIndex. ",error)
        finally:
            return index
            
    """ def createVectorIndex(self,client_service_name):
        index = None
        try:
            self.s3_fs.invalidate_cache()
            reader = SimpleDirectoryReader(
                input_dir=self.bucket_name+"/"+str(client_service_name),
                fs=self.s3_fs,
                recursive=True,  # recursively searches all subdirectories
            )
            documents = []
            for docs in reader.iter_data():
                for doc in docs:                    
                    doc.text = doc.text.upper()
                    documents.append(doc)
            
            #documents = reader.load_data()
            #index = VectorStoreIndex.from_documents(documents,transformations=[SentenceSplitter(chunk_size=self.sentence_chunk_size)])  
            index = VectorStoreIndex.from_documents(
                        documents=documents,
                        embed_model=self.embed_model,
                        llm=self.llm
                    )

            index.set_index_id(client_service_name)
            self.s3_client_obj.put_object(Bucket=self.index_s3_bucket_name,Body='', Key=str(client_service_name)+'/')
            index.storage_context.persist(persist_dir=self.index_s3_bucket_name+"/"+str(client_service_name), fs=self.s3_fs)
        except Exception as error:
            print("Exception in LLAMA_INDEX createVectorIndex. ",error)
        finally:
            return index """
        
    """ def getVectorIndex(self,vector_index_id):
        index = None
        try:
            index = load_index_from_storage(
                StorageContext.from_defaults(persist_dir=self.index_s3_bucket_name+"/"+str(vector_index_id), fs=self.s3_fs),
                index_id=vector_index_id,
            )
        except Exception as error:
            print("Exception in LLAMA_INDEX getVectorIndex. ",error)
        finally:
            return index """
        
    def getQueryEngine(self,index,similarity_top,collectionName,qa_prompt,transcript=""):
        query_engine = None
        try:
            qa_prompt_tmpl_str = (qa_prompt) 
            chat_history_str = json.dumps(transcript)
            qa_prompt_tmpl_str = qa_prompt_tmpl_str.replace("{{CHAT_HISTORY}}", chat_history_str)

            print(f"QA Prompt {qa_prompt_tmpl_str}")

            qa_prompt_tmpl = PromptTemplate(qa_prompt_tmpl_str, prompt_type=PromptType.QUESTION_ANSWER)

            filters = MetadataFilters(
                        filters=[ExactMatchFilter(key="role", value=str(collectionName))]
                    )
            retriever = VectorIndexRetriever(
                            index=index,
                            similarity_top_k=similarity_top,
                            filters=filters,
                        )     
            response_synthesizer = get_response_synthesizer(text_qa_template=qa_prompt_tmpl)

            if index:
                #query_engine = self.index.as_query_engine(similarity_top_k=5,response_mode="tree_summarize",streaming=True)
                #query_engine = index.as_query_engine(llm=self.llm,similarity_top_k=similarity_top,response_mode=self.QUERY_ENGINE_RESPONSE_MODE, verbose=True,filters=filters)                
                query_engine = RetrieverQueryEngine(retriever=retriever, response_synthesizer=response_synthesizer)
        except Exception as error:
            print("Exception in LLAMA_INDEX getQueryEngine. ",error)
        finally:
            return query_engine

    def askQuestions(self,query_engine,query,conversation_id):
        response = None
        sequence_id = 1
        values = []
        rag_response = None
        try:  
            # print(f"user query : {query}")
            response = query_engine.query(str(query))
            rag_response = response.response
            # print(f"Response {rag_response}")
            for node in response.source_nodes: 
                # print(f"ID : \n{node.node_id}")            
                # print(f"Score : \n{node.score}")
                # print(f"Text : \n{node.text}")     
                # print(f"Metadata : \n{node.metadata}")
                # print("\n\n")
                values.append((sequence_id,conversation_id,node.node_id,round(node.score,3),node.text,node.metadata,query,rag_response))                
                sequence_id = sequence_id+1
            self.insertRagResponseData(values)
        except Exception as error:
            print("Exception in LLAMA_INDEX createVectorIndex. ",error)
        finally:
            return response.response
